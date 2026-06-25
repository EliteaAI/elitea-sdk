"""
Pytest tests for EliteADocxMammothLoader.

These tests specifically verify:
1. Image handling with LLM (use_llm=True) properly invokes the LLM with image bytes
2. Image handling returns a placeholder transcript when no LLM is available
3. Image handling returns "Transcript is not available" when the LLM call fails
4. The fix for issue #4794 - ensuring image.open().read() is called correctly

Run:
  pytest tests/runtime/langchain/document_loaders/test_elitea_docx_mammoth_loader.py -v
"""

import io
from unittest.mock import MagicMock, patch, Mock
import pytest
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
    EliteADocxMammothLoader
)


class MockImageFile:
    """Mock for the file-like object returned by mammoth image.open()"""
    def __init__(self, data: bytes):
        self._data = data
        self._position = 0

    def read(self, size: int = -1) -> bytes:
        if size == -1:
            result = self._data[self._position:]
            self._position = len(self._data)
        else:
            result = self._data[self._position:self._position + size]
            self._position += size
        return result

    def __enter__(self):
        return self

    def __exit__(self, *args):
        pass


class MockMammothImage:
    """Mock for mammoth's Image object that has an open() method returning a file-like object"""
    def __init__(self, data: bytes, content_type: str = "image/png"):
        self._data = data
        self.content_type = content_type
        self.alt_text = None

    def open(self):
        return MockImageFile(self._data)


# Minimal valid 1x1 red pixel PNG (proper CRCs so python-docx can parse it)
MINIMAL_PNG_BYTES = (
    b'\x89PNG\r\n\x1a\n'
    b'\x00\x00\x00\rIHDR'
    b'\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00'
    b'\x90wS\xde'
    b'\x00\x00\x00\x0cIDAT'
    b'x\x9cc\xf8\xcf\xc0\x00\x00\x03\x01\x01\x00'
    b'\xc9\xfe\x92\xef'
    b'\x00\x00\x00\x00IEND\xaeB`\x82'
)


class TestDocxMammothLoaderImageHandling:
    """Test suite for image handling in EliteADocxMammothLoader"""

    def test_handle_image_with_llm_calls_read_on_file_object(self):
        """
        Test that __handle_image properly reads bytes from the file object.

        This is the core fix for issue #4794 - the image.open() returns a file-like
        object, and we must call .read() on it to get actual bytes before passing
        to perform_llm_prediction_for_image_bytes.
        """
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        # Create loader with mocked LLM
        loader = EliteADocxMammothLoader(file_path='/tmp/test.docx')
        mock_llm = MagicMock()
        mock_llm.invoke.return_value = MagicMock(content='LLM generated description')
        loader.llm = mock_llm
        loader.prompt = 'Describe this image'

        # Create mock mammoth image
        test_bytes = MINIMAL_PNG_BYTES
        mock_image = MockMammothImage(test_bytes)

        # Patch the perform_llm_prediction_for_image_bytes function to verify bytes are passed
        with patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
        ) as mock_predict:
            mock_predict.return_value = 'Mocked LLM description'

            # Call the handler
            result = loader._EliteADocxMammothLoader__handle_image(mock_image)

            # Verify the function was called with actual bytes (not file object)
            mock_predict.assert_called_once()
            call_args = mock_predict.call_args[0]

            # First argument should be bytes
            assert isinstance(call_args[0], bytes), \
                f"Expected bytes, got {type(call_args[0])}. This would fail before the fix!"
            assert call_args[0] == test_bytes, "Bytes should match the test image data"

            # Second argument should be the LLM
            assert call_args[1] == mock_llm

            # Third argument should be the prompt
            assert call_args[2] == 'Describe this image'

            # src is now a unique token; the payload (filename + transcript) is
            # stored verbatim in the payload map so parens/newlines survive.
            payload = loader._image_payload_map[result['src']]
            assert 'Mocked LLM description' in payload
            assert 'Image: image_1' in payload

    def test_handle_image_with_llm_returns_description(self):
        """Test that LLM-generated descriptions are returned correctly"""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        loader = EliteADocxMammothLoader(file_path='/tmp/test.docx')
        mock_llm = MagicMock()
        loader.llm = mock_llm
        loader.prompt = 'Describe this image'

        mock_image = MockMammothImage(MINIMAL_PNG_BYTES)

        with patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
        ) as mock_predict:
            mock_predict.return_value = 'A chart showing quarterly revenue growth'

            result = loader._EliteADocxMammothLoader__handle_image(mock_image)

            payload = loader._image_payload_map[result['src']]
            assert 'A chart showing quarterly revenue growth' in payload

    def test_handle_image_without_llm_returns_placeholder(self):
        """Without an LLM, no transcription backend exists → placeholder text."""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        loader = EliteADocxMammothLoader(file_path='/tmp/test.docx')
        loader.llm = None  # No LLM
        loader.prompt = None

        mock_image = MockMammothImage(MINIMAL_PNG_BYTES)

        result = loader._EliteADocxMammothLoader__handle_image(mock_image)

        # No LLM → no transcript available (Tesseract fallback removed)
        payload = loader._image_payload_map[result['src']]
        assert 'Transcript is not available' in payload
        assert 'Image: image_1' in payload

    def test_handle_image_fallback_on_llm_exception(self):
        """When the LLM call fails, the placeholder transcript is returned."""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        loader = EliteADocxMammothLoader(file_path='/tmp/test.docx')
        mock_llm = MagicMock()
        loader.llm = mock_llm
        loader.prompt = 'Describe this image'

        mock_image = MockMammothImage(MINIMAL_PNG_BYTES)

        with patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
        ) as mock_predict:
            mock_predict.side_effect = Exception('LLM service unavailable')

            result = loader._EliteADocxMammothLoader__handle_image(mock_image)

            # LLM failed and there is no OCR fallback → placeholder transcript
            payload = loader._image_payload_map[result['src']]
            assert 'Transcript is not available' in payload

    def test_handle_image_bug_reproduction_without_read(self):
        """
        Test that demonstrates the bug behavior - passing file object instead of bytes.

        This test verifies that bytes_to_base64 fails when given a file-like object,
        which was the root cause of issue #4794.
        """
        from elitea_sdk.runtime.langchain.tools.utils import bytes_to_base64

        # Create a file-like object (what image.open() returns)
        file_obj = io.BytesIO(b'test image data')

        # This should fail - demonstrating the bug
        with pytest.raises(TypeError, match="bytes-like object"):
            bytes_to_base64(file_obj)

        # The correct approach is to call .read() first
        file_obj.seek(0)
        actual_bytes = file_obj.read()
        result = bytes_to_base64(actual_bytes)
        assert result == 'dGVzdCBpbWFnZSBkYXRh'  # base64 of 'test image data'


class TestDocxMammothLoaderIntegration:
    """Integration tests for EliteADocxMammothLoader"""

    def test_loader_initialization(self):
        """Test that loader initializes correctly with various kwargs"""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        # Test with file_path
        loader = EliteADocxMammothLoader(file_path='/tmp/test.docx')
        assert loader.path == '/tmp/test.docx'
        assert loader.llm is None
        assert loader.prompt is None
        assert loader.max_tokens == 512

        # Test with all kwargs
        mock_llm = MagicMock()
        loader = EliteADocxMammothLoader(
            file_path='/tmp/test.docx',
            llm=mock_llm,
            prompt='Custom prompt',
            max_tokens=1024,
            extract_images=True
        )
        assert loader.llm == mock_llm
        assert loader.prompt == 'Custom prompt'
        assert loader.max_tokens == 1024
        assert loader.extract_images is True

    def test_loader_with_file_content(self):
        """Test that loader can be initialized with in-memory content"""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        loader = EliteADocxMammothLoader(
            file_content=b'docx content',
            file_name='test.docx'
        )
        assert loader.file_content == b'docx content'
        assert loader.file_name == 'test.docx'


class TestProcessContentByTypeDocxWithLLM:
    """Test process_content_by_type with .docx files and use_llm=True"""

    def test_use_llm_config_is_properly_passed(self):
        """
        Test that use_llm=True in chunking_config properly enables LLM for docx processing.

        This tests the full flow from process_content_by_type through to the loader.
        """
        from elitea_sdk.tools.utils.content_parser import process_content_by_type
        from elitea_sdk.runtime.langchain.document_loaders.constants import LoaderProperties

        # Minimal valid DOCX is complex to create, so we'll mock the loader
        with patch(
            'elitea_sdk.tools.utils.content_parser.loaders_map'
        ) as mock_loaders_map:
            # Setup mock loader config
            mock_loader_cls = MagicMock()
            mock_loader_instance = MagicMock()
            mock_loader_cls.return_value = mock_loader_instance
            mock_loader_instance.load.return_value = iter([])

            mock_loaders_map.get.return_value = {
                'class': mock_loader_cls,
                'kwargs': {'extract_images': True},
                'is_multimodal_processing': True,
                'allowed_to_override': {
                    LoaderProperties.LLM.value: False,
                    LoaderProperties.PROMPT_DEFAULT.value: False,
                    LoaderProperties.PROMPT.value: "",
                    'mode': 'paged',
                    'max_tokens': 512,
                }
            }

            # Create mock LLM
            mock_llm = MagicMock()

            # Chunking config with use_llm=True
            chunking_config = {
                '.docx': {
                    'max_tokens': 512,
                    'mode': 'paged',
                    'prompt': '',
                    'use_default_prompt': False,
                    'use_llm': True
                }
            }

            # Process content
            list(process_content_by_type(
                b'fake docx content',
                'test.docx',
                llm=mock_llm,
                chunking_config=chunking_config
            ))

            # Verify loader was called with llm parameter
            mock_loader_cls.assert_called_once()
            call_kwargs = mock_loader_cls.call_args[1]

            # The llm should be passed to the loader
            assert 'llm' in call_kwargs, "LLM should be passed to loader when use_llm=True"
            assert call_kwargs['llm'] == mock_llm


# ---------------------------------------------------------------------------
# EL-4629 — LLM-only image handler & get_file_metadata
# ---------------------------------------------------------------------------

def _make_docx_with_image():
    """Build a minimal DOCX (bytes) with one paragraph and one embedded PNG image."""
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_paragraph("Hello, this document has an image.")
    doc.add_picture(io.BytesIO(MINIMAL_PNG_BYTES))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxGetFileMetadata:
    """Tests for EliteADocxMammothLoader.get_file_metadata (EL-4629)."""

    def test_metadata_reports_image_count(self):
        """File with 1 embedded image should return image_count=1."""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        docx_bytes = _make_docx_with_image()
        meta = EliteADocxMammothLoader.get_file_metadata(
            filename='report.docx', file_content=docx_bytes, file_size=len(docx_bytes))

        assert meta['image_count'] >= 1
        assert len(meta['image_names']) >= 1
        instr = meta['instruction_for_readFile']
        assert 'is_capture_image' in instr['first_class_params']
        assert 'extracted_images_names' in instr['extra_params']
        assert 'prompt' in instr['extra_params']

    def test_metadata_no_content_returns_zero_count(self):
        """Without file_content, image_count should be 0."""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        meta = EliteADocxMammothLoader.get_file_metadata(
            filename='doc.docx', file_content=None)

        assert meta['image_count'] == 0
        assert meta['image_names'] == []
        assert 'extracted_images_names' in meta['instruction_for_readFile']['extra_params']

    def test_metadata_no_images_document(self):
        """A plain text DOCX with no images should return image_count=0."""
        from docx import Document as DocxDocument
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        doc = DocxDocument()
        doc.add_paragraph("Just text, no images.")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        meta = EliteADocxMammothLoader.get_file_metadata(
            filename='plain.docx', file_content=docx_bytes)

        assert meta['image_count'] == 0
        assert meta['image_names'] == []


# ---------------------------------------------------------------------------
# v3 — Placeholders, dedup cache, selective image reading
# ---------------------------------------------------------------------------

def _make_docx_with_named_images(image_count=2):
    """Build a DOCX with multiple distinct embedded images."""
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_paragraph("Document with multiple images.")
    for i in range(image_count):
        doc.add_paragraph(f"Before image {i + 1}.")
        doc.add_picture(io.BytesIO(MINIMAL_PNG_BYTES))
        doc.add_paragraph(f"After image {i + 1}.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxImagePlaceholders:
    """Tests for image placeholder markers when extract_images=False."""

    def test_placeholder_contains_image_filename(self):
        """Non-image read should show placeholder with the media filename."""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        docx_bytes = _make_docx_with_image()
        loader = EliteADocxMammothLoader(
            file_content=docx_bytes, file_name='test.docx',
            extract_images=False)

        content = loader.get_content()

        # Should contain the placeholder with the actual image filename
        assert 'Image:' in content
        assert 'image1.png' in content
        assert 'get_file_metadata' in content

    def test_placeholder_does_not_call_llm(self):
        """Placeholder mode should never invoke the LLM."""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        docx_bytes = _make_docx_with_image()
        mock_llm = MagicMock()
        loader = EliteADocxMammothLoader(
            file_content=docx_bytes, file_name='test.docx',
            llm=mock_llm, extract_images=False)

        with patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
        ) as mock_predict:
            loader.get_content()

        mock_predict.assert_not_called()


class TestDocxImageDedup:
    """Tests for image deduplication cache when extract_images=True."""

    def test_same_image_processed_once(self):
        """Same image referenced multiple times → LLM called only once."""
        from docx import Document as DocxDocument
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        # Build DOCX referencing the same image twice
        doc = DocxDocument()
        doc.add_paragraph("First reference.")
        doc.add_picture(io.BytesIO(MINIMAL_PNG_BYTES))
        doc.add_paragraph("Second reference.")
        doc.add_picture(io.BytesIO(MINIMAL_PNG_BYTES))
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        mock_llm = MagicMock()
        loader = EliteADocxMammothLoader(
            file_content=docx_bytes, file_name='test.docx',
            llm=mock_llm, extract_images=True)

        with patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
        ) as mock_predict:
            mock_predict.return_value = 'A red pixel'
            content = loader.get_content()

        # Same PNG bytes → should be called only once (dedup)
        assert mock_predict.call_count == 1
        # Second occurrence should say "already transcribed"
        assert 'already transcribed' in content

    def test_transcript_format_includes_filename(self):
        """First occurrence should include 'Image: filename, transcript'."""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        docx_bytes = _make_docx_with_image()
        mock_llm = MagicMock()
        loader = EliteADocxMammothLoader(
            file_content=docx_bytes, file_name='test.docx',
            llm=mock_llm, extract_images=True)

        with patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
        ) as mock_predict:
            mock_predict.return_value = 'A bar chart showing revenue'
            content = loader.get_content()

        assert 'image1.png' in content
        assert 'A bar chart showing revenue' in content


class TestDocxSelectiveImageReading:
    """Tests for read_images_only + extracted_images_names."""

    def test_read_images_only_returns_json(self):
        """read_images_only=True returns JSON dict of image transcripts."""
        import json as json_mod
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        docx_bytes = _make_docx_with_image()
        mock_llm = MagicMock()

        # Get the actual image filename from metadata
        meta = EliteADocxMammothLoader.get_file_metadata(
            filename='test.docx', file_content=docx_bytes)
        img_name = meta['image_names'][0]

        loader = EliteADocxMammothLoader(
            file_content=docx_bytes, file_name='test.docx',
            llm=mock_llm, extract_images=False,
            extracted_images_names=[img_name],
            read_images_only=True)

        with patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
        ) as mock_predict:
            mock_predict.return_value = 'Described image content'
            result = loader.get_content()

        parsed = json_mod.loads(result)
        assert img_name in parsed
        assert parsed[img_name] == 'Described image content'
        mock_predict.assert_called_once()

    def test_read_images_only_nonexistent(self):
        """Requesting a non-existent image returns error in the dict."""
        import json as json_mod
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        docx_bytes = _make_docx_with_image()
        loader = EliteADocxMammothLoader(
            file_content=docx_bytes, file_name='test.docx',
            llm=MagicMock(), extract_images=False,
            extracted_images_names=['nonexistent.png'],
            read_images_only=True)

        result = loader.get_content()
        parsed = json_mod.loads(result)
        assert 'nonexistent.png' in parsed
        assert 'not found' in parsed['nonexistent.png']

    def test_read_images_only_skips_mammoth(self):
        """read_images_only should not call mammoth convert_to_html."""
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        docx_bytes = _make_docx_with_image()
        meta = EliteADocxMammothLoader.get_file_metadata(
            filename='test.docx', file_content=docx_bytes)
        img_name = meta['image_names'][0]

        loader = EliteADocxMammothLoader(
            file_content=docx_bytes, file_name='test.docx',
            llm=MagicMock(), extract_images=False,
            extracted_images_names=[img_name],
            read_images_only=True)

        with patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.convert_to_html'
        ) as mock_convert, \
             patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
        ) as mock_predict:
            mock_predict.return_value = 'transcript'
            loader.get_content()

        mock_convert.assert_not_called()


class TestDocxScanImageReferences:
    """Tests for _scan_image_references pre-scan."""

    def test_scan_finds_image_names(self):
        """Pre-scan should find embedded image filenames in document order."""
        from docx import Document as DocxDocument
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        docx_bytes = _make_docx_with_named_images(image_count=2)
        doc = DocxDocument(io.BytesIO(docx_bytes))

        loader = EliteADocxMammothLoader(
            file_content=docx_bytes, file_name='test.docx')
        loader._scan_image_references(doc)

        assert len(loader._image_ref_order) == 2
        # python-docx names them image1.png, image2.png
        assert all(name.endswith('.png') for name in loader._image_ref_order)

    def test_scan_empty_document(self):
        """Document without images → empty reference list."""
        from docx import Document as DocxDocument
        from elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader import (
            EliteADocxMammothLoader
        )

        doc = DocxDocument()
        doc.add_paragraph("No images here.")
        buf = io.BytesIO()
        doc.save(buf)

        loader = EliteADocxMammothLoader(
            file_content=buf.getvalue(), file_name='test.docx')
        loader._scan_image_references(DocxDocument(io.BytesIO(buf.getvalue())))

        assert loader._image_ref_order == []


class TestDocxHoistImagesFromHeadings:
    """Tests for __hoist_images_from_headings (issue #5333)."""

    def test_image_inside_heading_is_hoisted_into_following_paragraph(self):
        """An <img> inside a heading must be relocated into its own <p> right
        after the heading so markdownify keeps it; the heading keeps its text
        and the image src (which carries the transcript) is preserved.
        """
        loader = EliteADocxMammothLoader(file_path='/tmp/test.docx')

        html = (
            '<h2>Section title<img src="Image: image1.png, a bubble sort diagram" /></h2>'
            '<p>Body text.</p>'
        )

        result = loader._EliteADocxMammothLoader__hoist_images_from_headings(html)
        soup = BeautifulSoup(result, 'html.parser')

        heading = soup.find('h2')
        # Heading keeps its text but no longer holds the image.
        assert heading.get_text(strip=True) == 'Section title'
        assert heading.find('img') is None

        # The image now lives in a <p> immediately after the heading, src intact.
        hoisted = heading.find_next_sibling()
        assert hoisted.name == 'p'
        img = hoisted.find('img')
        assert img is not None
        assert img['src'] == 'Image: image1.png, a bubble sort diagram'

    def test_image_only_heading_is_dropped_not_left_empty(self):
        """A heading that held only the image must be removed, otherwise it
        markdownifies to a dangling empty '#' header line."""
        loader = EliteADocxMammothLoader(file_path='/tmp/test.docx')

        html = '<h1><img src="Image: image1.png, diagram" /></h1><p>Body.</p>'

        result = loader._EliteADocxMammothLoader__hoist_images_from_headings(html)
        soup = BeautifulSoup(result, 'html.parser')

        assert soup.find('h1') is None  # empty heading dropped
        assert soup.find('img') is not None  # image preserved in its new <p>


class TestDocxImageTranscriptPreservation:
    """Transcripts with parentheses/newlines survive postprocessing (#5333).

    Regression for the previous approach that round-tripped the transcript
    through the markdown image URL and truncated it at the first ')'.
    """

    def test_postprocess_preserves_parentheses_and_newlines(self):
        loader = EliteADocxMammothLoader(file_path='/tmp/test.docx')

        transcript = ("Bubble sort (ascending) step 1.\n"
                      "Trend: increasing (left to right).\nNumbers: (5, 3, 8).")
        token = loader._EliteADocxMammothLoader__register_image_payload(
            f"Image: image1.png, {transcript}")

        # markdownify renders the token-src image as ![](token)
        md = f"# Title\n\n![]({token})\n\nBody text."
        result = loader._EliteADocxMammothLoader__postprocess_original_md(md)

        assert transcript in result  # verbatim — no truncation at ')'
        assert token not in result  # token resolved
        assert "**Image Transcript:**" in result


class TestDocxLoaderReuse:
    """Per-conversion image state must not leak across reuse (#5333)."""

    def test_reset_clears_dedup_cache_so_repeated_image_re_transcribes(self):
        """A reused instance processing a second document whose image bytes
        match the first must NOT emit a stale 'already transcribed' back-ref;
        it must re-transcribe the image in the new document."""
        loader = EliteADocxMammothLoader(file_path='/tmp/test.docx')
        loader.llm = MagicMock()
        loader.prompt = 'Describe'

        with patch(
            'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader'
            '.perform_llm_prediction_for_image_bytes'
        ) as mock_predict:
            mock_predict.return_value = 'A unique description'

            # First conversion
            loader._reset_image_state()
            r1 = loader._EliteADocxMammothLoader__handle_image(
                MockMammothImage(MINIMAL_PNG_BYTES))
            assert 'A unique description' in loader._image_payload_map[r1['src']]

            # Second conversion, SAME image bytes — reset must drop the dedup hit
            loader._reset_image_state()
            r2 = loader._EliteADocxMammothLoader__handle_image(
                MockMammothImage(MINIMAL_PNG_BYTES))
            payload2 = loader._image_payload_map[r2['src']]
            assert 'already transcribed' not in payload2  # no stale back-ref
            assert 'A unique description' in payload2  # re-transcribed

    def test_convert_docx_to_markdown_resets_image_state(self):
        """Guard the call site: the conversion entry point must invoke
        _reset_image_state, so deleting that one line is caught by a test
        (not only the helper's own behavior)."""
        buf = io.BytesIO()
        document = DocxDocument()
        document.add_paragraph("hello world")
        document.save(buf)

        loader = EliteADocxMammothLoader(
            file_content=buf.getvalue(), file_name='t.docx')
        with patch.object(
            loader, '_reset_image_state', wraps=loader._reset_image_state
        ) as spy:
            loader._convert_docx_to_markdown(io.BytesIO(buf.getvalue()))
        spy.assert_called_once()
