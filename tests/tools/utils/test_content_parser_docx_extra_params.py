"""Integration test for extra_params plumbing through parse_file_content for DOCX (EL-4629).

Verifies that when extra_params={"extract_images": True} is passed to
parse_file_content for a DOCX file, the EliteADocxMammothLoader receives
extract_images=True and _llm_only_images=True, invoking the LLM-only
image handler.
"""

import io

import pytest
from unittest.mock import MagicMock, patch

from elitea_sdk.tools.utils.content_parser import parse_file_content


# Minimal valid 1x1 red pixel PNG (proper CRCs so python-docx can parse it)
MINIMAL_PNG_BYTES = (
    b'\x89PNG\r\n\x1a\n'
    b'\x00\x00\x00\rIHDR'
    b'\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00'
    b'\x90wS\xde'
    b'\x00\x00\x00\x0cIDAT'
    b'x\x9cc\xf8\xcf\xc0\x00\x00\x03\x01\x01\x00'
    b'\xc9\xfe\x92\xef'
    b'\x00\x00\x00\x00IEND\xaeB`\x82'
)


def _make_docx_with_image():
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_paragraph("Document with an image.")
    doc.add_picture(io.BytesIO(MINIMAL_PNG_BYTES))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parse_file_content_docx_extract_images_via_extra_params():
    """extra_params={"extract_images": True} should trigger LLM image transcription."""
    docx_bytes = _make_docx_with_image()
    mock_llm = MagicMock()

    with patch(
        'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
    ) as mock_predict:
        mock_predict.return_value = 'TRANSCRIBED-VIA-EXTRA-PARAMS'

        result = parse_file_content(
            file_name='report.docx',
            file_content=docx_bytes,
            llm=mock_llm,
            extra_params={"extract_images": True},
        )

    assert isinstance(result, str)
    assert 'TRANSCRIBED-VIA-EXTRA-PARAMS' in result
    mock_predict.assert_called()


def test_parse_file_content_docx_without_extra_params_no_transcription():
    """Without extra_params, DOCX images should NOT be transcribed."""
    docx_bytes = _make_docx_with_image()
    mock_llm = MagicMock()

    with patch(
        'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
    ) as mock_predict:
        result = parse_file_content(
            file_name='report.docx',
            file_content=docx_bytes,
            llm=mock_llm,
        )

    mock_predict.assert_not_called()
    assert isinstance(result, str)


def test_parse_file_content_docx_extract_images_false_explicit():
    """extra_params={"extract_images": False} should explicitly disable transcription."""
    docx_bytes = _make_docx_with_image()
    mock_llm = MagicMock()

    with patch(
        'elitea_sdk.runtime.langchain.document_loaders.EliteADocxMammothLoader.perform_llm_prediction_for_image_bytes'
    ) as mock_predict:
        result = parse_file_content(
            file_name='report.docx',
            file_content=docx_bytes,
            llm=mock_llm,
            extra_params={"extract_images": False},
        )

    mock_predict.assert_not_called()
