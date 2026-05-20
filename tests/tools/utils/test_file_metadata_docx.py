"""Tests for file_metadata registry with DOCX files (EL-4629).

Covers:
  * DOCX handler advertises image_count + instruction_for_readFile
  * Generic fallback structure remains valid
  * image_count reflects actual embedded images
"""

import io

import pytest

from elitea_sdk.tools.utils.file_metadata import get_file_metadata


# Minimal valid 1x1 red pixel PNG (proper CRCs so python-docx can parse it)
MINIMAL_PNG_BYTES = (
    b'\x89PNG\r\n\x1a\n'
    b'\x00\x00\x00\rIHDR'
    b'\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00'
    b'\x90wS\xde'
    b'\x00\x00\x00\x0cIDAT'
    b'x\x9cc\xf8\xcf\xc0\x00\x00\x03\x01\x01\x00'
    b'\xc9\xfe\x92\xef'
    b'\x00\x00\x00\x00IEND\xaeB`\x82'
)


def _make_docx_bytes(with_image=True):
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_paragraph("Test document.")
    if with_image:
        doc.add_picture(io.BytesIO(MINIMAL_PNG_BYTES))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_metadata_returns_image_count_and_instruction():
    data = _make_docx_bytes(with_image=True)

    meta = get_file_metadata("report.docx", file_content=data, file_size=len(data))

    assert meta["filename"] == "report.docx"
    assert meta["extension"] == ".docx"
    assert meta["filesize"] == len(data)
    assert meta["image_count"] >= 1
    assert len(meta["image_names"]) >= 1
    instr = meta["instruction_for_readFile"]
    assert "is_capture_image" in instr["first_class_params"]
    assert "extracted_images_names" in instr["extra_params"]
    assert "prompt" in instr["extra_params"]
    assert "JSON string" in instr["notes"]


def test_docx_metadata_no_images():
    data = _make_docx_bytes(with_image=False)

    meta = get_file_metadata("plain.docx", file_content=data, file_size=len(data))

    assert meta["image_count"] == 0
    assert meta["image_names"] == []
    assert "extracted_images_names" in meta["instruction_for_readFile"]["extra_params"]


def test_docx_metadata_no_content():
    meta = get_file_metadata("doc.docx", file_content=None, file_size=4096)

    assert meta["image_count"] == 0
    assert meta["image_names"] == []
    assert meta["filesize"] == 4096
    assert "extracted_images_names" in meta["instruction_for_readFile"]["extra_params"]
