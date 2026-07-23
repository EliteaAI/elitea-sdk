import hashlib
import json
import re
import uuid
from io import BytesIO

import mammoth.images
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from mammoth import convert_to_html
from markdownify import markdownify
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from bs4 import BeautifulSoup

from elitea_sdk.tools.chunkers.sematic.markdown_chunker import markdown_by_headers_chunker
from .utils import perform_llm_prediction_for_image_bytes


class EliteADocxMammothLoader(BaseLoader):
    """
    Loader for Docx files using Mammoth to convert to HTML, with image handling,
    and then Markdownify to convert HTML to markdown.
    Detects bordered paragraphs and text boxes and treats them as code blocks.
    """
    @classmethod
    def get_file_metadata(cls, *, filename: str,
                          file_content=None,
                          file_size=None) -> dict:
        import logging
        _logger = logging.getLogger(__name__)
        image_count = 0
        image_names = []
        total_lines = 0
        if file_content:
            try:
                doc = DocxDocument(BytesIO(file_content) if isinstance(file_content, (bytes, bytearray)) else file_content)
                for rel in doc.part.rels.values():
                    try:
                        tp = rel.target_part
                    except Exception:  # pylint: disable=broad-except
                        continue
                    if hasattr(tp, 'content_type') and tp.content_type.startswith('image/'):
                        image_count += 1
                        image_names.append(tp.partname.filename)
            except Exception as exc:  # pylint: disable=broad-except
                _logger.warning("Failed to inspect %s: %s", filename, exc)

            # A chunked read renders the whole docx to markdown anyway (mammoth
            # has no bounded-slice API), so count lines on that same rendered
            # text. extract_images=False keeps it deterministic and LLM-free.
            try:
                rendered = cls(file_content=file_content, file_name=filename,
                               extract_images=False).get_content()
                total_lines = len(rendered.splitlines())
            except Exception as exc:  # pylint: disable=broad-except
                _logger.warning("Failed to render %s for line count: %s", filename, exc)

        range_hint = f"Valid range 1..{total_lines}. " if total_lines else ""
        instruction = {
            "first_class_params": {
                "is_capture_image": (
                    "set to true to transcribe ALL embedded images "
                    "via the AI vision model"
                ),
                "start_line": (
                    "integer (1-indexed, inclusive) — first line of the "
                    f"converted markdown to return. {range_hint}"
                    "Omit to read from the start."
                ),
                "end_line": (
                    "integer (1-indexed, inclusive) — last line of the "
                    f"converted markdown to return. {range_hint}"
                    "Omit to read to the end."
                ),
            },
            "extra_params": {
                "extracted_images_names": (
                    "list of image filenames to selectively transcribe "
                    '(e.g. ["image1.png", "image2.png"])'
                ),
                "read_images_only": (
                    "boolean — when true together with "
                    "extracted_images_names, returns only image "
                    "transcripts without the document text"
                ),
                "prompt": (
                    "string (optional) — custom prompt for the vision "
                    "model. If omitted, the default image-processing "
                    "prompt is used."
                ),
            },
            "notes": (
                f"Document renders to {total_lines} markdown lines. Use "
                "start_line/end_line to page through it in bounded chunks. "
                "total_lines counts the text WITHOUT image transcripts; passing "
                "is_capture_image adds transcript lines and shifts the count. "
                "Image transcription adds latency proportional to image count. "
                "Pass extra_params as a JSON string."
            ),
        }
        return {
            "unit": "lines",
            "total_lines": total_lines,
            "image_count": image_count,
            "image_names": image_names,
            "instruction_for_readFile": instruction,
        }

    def __init__(self, **kwargs):
        """Initializes EliteADocxMammothLoader."""
        self.path =  kwargs.get('file_path')
        self.file_content = kwargs.get('file_content')
        self.file_name = kwargs.get('file_name')
        self.extract_images = kwargs.get('extract_images')
        self.llm = kwargs.get("llm")
        self.prompt = kwargs.get("prompt")
        self.max_tokens = kwargs.get('max_tokens', 512)
        self.extracted_images_names = kwargs.get('extracted_images_names')
        self.read_images_only = kwargs.get('read_images_only', False)
        self.image_cache = kwargs.get('image_cache')
        # Per-conversion dedup: MD5(image_bytes) → transcript string. Used only
        # to emit "[already transcribed, see above]" back-references for
        # duplicate bytes inside a single document — distinct from the shared
        # cross-invocation ``image_cache`` (ImageDescriptionCache) above.
        self._per_doc_transcripts = {}
        # Ordered list of image filenames as they appear in the document
        self._image_ref_order = []
        # Counter to track current position during mammoth callbacks
        self._image_ref_index = 0
        # token → image payload text (filename + transcript). Image handlers
        # embed the token as the <img> src so the transcript — which routinely
        # contains parentheses/newlines (LLM vision output) — is never parsed
        # as part of the markdown image URL. Resolved in __postprocess_original_md.
        self._image_payload_map = {}

    def _reset_image_state(self):
        """Clear per-conversion image caches.

        Reset together so they cannot drift apart: the token→payload map and
        the MD5→transcript dedup cache. Without clearing the dedup cache, a
        reused instance processing a second document whose image bytes match
        one from the first would short-circuit to a "[already transcribed,
        see above]" back-reference that points at nothing — and lose that
        image's transcript. (``_image_ref_*`` are reset by
        ``_scan_image_references`` on every conversion.)
        """
        self._image_payload_map = {}
        self._per_doc_transcripts = {}

    def __register_image_payload(self, payload_text: str) -> str:
        """Register an image payload and return a unique, paren/space-free token.

        The token is embedded as the <img> ``src``; markdownify turns it into
        ``![](<token>)`` and __postprocess_original_md swaps it back for the
        payload verbatim. This avoids corrupting transcripts that contain ``)``
        or newlines, which the previous ``src``-as-text approach truncated.
        """
        token = f"eliteaimgtoken{uuid.uuid4().hex}"
        self._image_payload_map[token] = payload_text
        return token

    def _scan_image_references(self, doc):
        """Pre-scan document XML for image references in document order.

        Builds an ordered list of image filenames (from ``word/media/``)
        matching the order mammoth will invoke its ``convert_image`` callback.
        Also builds a mapping from relationship-ID to filename for direct
        image extraction.
        """
        # Build rId → filename mapping from relationships
        self._rid_to_filename = {}
        self._filename_to_rel = {}
        for rel in doc.part.rels.values():
            try:
                tp = rel.target_part
            except Exception:  # pylint: disable=broad-except
                continue  # skip External relationships (hyperlinks etc.)
            if hasattr(tp, 'content_type') and tp.content_type.startswith('image/'):
                fname = tp.partname.filename
                self._rid_to_filename[rel.rId] = fname
                self._filename_to_rel[fname] = rel

        # Walk body XML in document order to find <a:blip r:embed="rIdN">
        ns = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }
        self._image_ref_order = []
        for blip in doc.element.body.iter(qn('a:blip')):
            rid = blip.get(qn('r:embed'))
            if rid and rid in self._rid_to_filename:
                self._image_ref_order.append(self._rid_to_filename[rid])

        self._image_ref_index = 0

    def _get_current_image_name(self):
        """Return the filename for the current image callback invocation."""
        if self._image_ref_index < len(self._image_ref_order):
            name = self._image_ref_order[self._image_ref_index]
        else:
            name = f"image_{self._image_ref_index + 1}"
        self._image_ref_index += 1
        return name

    def __handle_image(self, image) -> dict:
        """Handles image processing with dedup cache.

        On first encounter of unique image bytes: transcribe via the LLM
        vision model, cache result, return transcript with filename.  On
        duplicate bytes: return back-reference placeholder.
        """
        from elitea_sdk.tools.utils.content_parser import image_processing_prompt

        image_name = self._get_current_image_name()
        try:
            with image.open() as image_file:
                image_bytes = image_file.read()

            img_hash = hashlib.md5(image_bytes).hexdigest()

            # Check per-conversion dedup
            if img_hash in self._per_doc_transcripts:
                return {"src": self.__register_image_payload(
                    f"Image: {image_name} [already transcribed, see above]"),
                    "alt": image_name}

            # Process image
            transcript = None
            if self.llm:
                try:
                    source = self.path or self.file_name or "docx"
                    transcript = perform_llm_prediction_for_image_bytes(
                        image_bytes, self.llm,
                        self.prompt if self.prompt else image_processing_prompt,
                        cache=self.image_cache,
                        image_name=source,
                    )
                except Exception:
                    transcript = None

            if transcript is None:
                transcript = "Transcript is not available"

            self._per_doc_transcripts[img_hash] = transcript
            return {"src": self.__register_image_payload(
                f"Image: {image_name}, {transcript}"), "alt": image_name}

        except Exception:
            return {"src": self.__register_image_payload(
                f"Image: {image_name}, Transcript is not available"),
                "alt": image_name}

    def __placeholder_image_handler(self, image) -> dict:
        """Placeholder handler for non-image reads.

        Returns a marker with the image filename so the LLM knows an image
        exists and can request selective expansion later.
        """
        image_name = self._get_current_image_name()
        return {
            "src": self.__register_image_payload(
                f"Image: {image_name}, you can selectively read it, "
                "call get_file_metadata to figure out how"
            ),
            "alt": image_name,
        }

    def __default_image_handler(self, image) -> dict:
        """Default image handler: returns a placeholder string."""
        return {"src": "Transcript is not available"}


    def __postprocess_original_md(self, original_md: str) -> str:
        """Swap image-payload tokens back for formatted transcript blocks.

        Image handlers embed a unique token as the <img> ``src``; markdownify
        renders it as ``![](<token>)``. Each token is replaced with its payload
        taken verbatim from ``_image_payload_map`` — so parentheses and newlines
        in the transcript (typical of LLM vision output) survive intact, unlike
        the previous approach that parsed the transcript out of the markdown URL.
        """
        if not self._image_payload_map:
            return original_md

        # ![alt](TOKEN) or ![alt](TOKEN "title") → transcript block
        image_md_pattern = re.compile(
            r'!\[[^\]]*\]\((eliteaimgtoken[0-9a-f]+)(?:\s+"[^"]*")?\)')

        def replace_token(match):
            payload = self._image_payload_map.get(match.group(1))
            if payload is None:
                return match.group(0)
            return f"\n**Image Transcript:**\n{payload}\n"

        new_md = image_md_pattern.sub(replace_token, original_md)

        # Defensive: if markdownify emitted a token outside image syntax
        # (e.g. as bare text), still swap it for its transcript.
        for token, payload in self._image_payload_map.items():
            if token in new_md:
                new_md = new_md.replace(
                    token, f"\n**Image Transcript:**\n{payload}\n")

        return new_md

    def __has_border(self, paragraph):
        """
        Check if a paragraph has border formatting.
        
        Args:
            paragraph: A python-docx Paragraph object.
            
        Returns:
            bool: True if paragraph has any border, False otherwise.
        """
        pPr = paragraph._element.pPr
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                # Check if any border side exists (top, bottom, left, right)
                for side in ['top', 'bottom', 'left', 'right']:
                    border = pBdr.find(qn(f'w:{side}'))
                    if border is not None:
                        # Check if border is not "none" or has a width
                        val = border.get(qn('w:val'))
                        if val and val != 'none':
                            return True
        return False

    def __find_text_boxes(self, doc):
        """
        Find all text boxes in document by searching OOXML structure.
        Text boxes are typically in w:txbxContent elements.
        
        Args:
            doc: A python-docx Document object.
            
        Returns:
            list: List of tuples (element, paragraphs_inside_textbox).
        """
        text_boxes = []
        
        # Iterate through document body XML to find text box content elements
        for element in doc.element.body.iter():
            # Look for text box content elements
            if element.tag.endswith('txbxContent'):
                # Collect all paragraphs inside this text box
                txbx_paragraphs = []
                for txbx_para_element in element.iter():
                    if txbx_para_element.tag.endswith('p'):
                        txbx_paragraphs.append(txbx_para_element)
                
                if txbx_paragraphs:
                    text_boxes.append((element, txbx_paragraphs))
        
        return text_boxes

    def __create_marker_paragraph(self, marker_text):
        """
        Create a paragraph element with marker text.
        
        Args:
            marker_text (str): The marker text to insert.
            
        Returns:
            Element: An OOXML paragraph element.
        """
        from docx.oxml import OxmlElement
        
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = marker_text
        r.append(t)
        p.append(r)
        return p

    def __inject_markers_for_paragraph(self, paragraph, start_marker, end_marker):
        """
        Inject marker paragraphs before and after a bordered paragraph.
        
        Args:
            paragraph: A python-docx Paragraph object.
            start_marker (str): The start marker text.
            end_marker (str): The end marker text.
        """
        # Insert start marker paragraph before
        marker_p_start = self.__create_marker_paragraph(start_marker)
        paragraph._element.addprevious(marker_p_start)
        
        # Insert end marker paragraph after
        marker_p_end = self.__create_marker_paragraph(end_marker)
        paragraph._element.addnext(marker_p_end)

    def __inject_markers_for_textbox(self, textbox_element, paragraph_elements, start_marker, end_marker):
        """
        Inject markers around text box content.
        
        Args:
            textbox_element: The w:txbxContent element.
            paragraph_elements: List of paragraph elements inside the text box.
            start_marker (str): The start marker text.
            end_marker (str): The end marker text.
        """
        if not paragraph_elements:
            return
        
        # Insert start marker before first paragraph in text box
        first_para = paragraph_elements[0]
        marker_p_start = self.__create_marker_paragraph(start_marker)
        first_para.addprevious(marker_p_start)
        
        # Insert end marker after last paragraph in text box
        last_para = paragraph_elements[-1]
        marker_p_end = self.__create_marker_paragraph(end_marker)
        last_para.addnext(marker_p_end)

    def __detect_and_mark_bordered_content(self, docx_stream):
        """
        Detects bordered paragraphs and text boxes, injects unique markers around them.
        Groups consecutive bordered paragraphs into single code blocks.
        
        Args:
            docx_stream: A file-like object containing the DOCX document.
            
        Returns:
            tuple: (modified_docx_stream, start_marker, end_marker)
        """
        # Load document with python-docx
        doc = DocxDocument(docx_stream)
        
        # Generate unique markers to avoid conflicts with document content
        unique_id = uuid.uuid4().hex[:8]
        start_marker = f"<<<BORDERED_BLOCK_START_{unique_id}>>>"
        end_marker = f"<<<BORDERED_BLOCK_END_{unique_id}>>>"
        
        # Group consecutive bordered paragraphs together
        bordered_groups = []
        current_group = []
        
        for para in doc.paragraphs:
            if self.__has_border(para):
                current_group.append(para)
            else:
                if current_group:
                    # End of a bordered group
                    bordered_groups.append(current_group)
                    current_group = []
        
        # Don't forget the last group if document ends with bordered paragraphs
        if current_group:
            bordered_groups.append(current_group)
        
        # Collect all text boxes
        # text_boxes = self.__find_text_boxes(doc)
        
        # Inject markers around each group of consecutive bordered paragraphs
        for group in bordered_groups:
            if group:
                # Add start marker before first paragraph in group
                first_para = group[0]
                marker_p_start = self.__create_marker_paragraph(start_marker)
                first_para._element.addprevious(marker_p_start)
                
                # Add end marker after last paragraph in group
                last_para = group[-1]
                marker_p_end = self.__create_marker_paragraph(end_marker)
                last_para._element.addnext(marker_p_end)
        
        # Inject markers around text box content
        # for textbox_element, para_elements in text_boxes:
        #     self.__inject_markers_for_textbox(textbox_element, para_elements, start_marker, end_marker)
        
        # Save modified document to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output, start_marker, end_marker

    def __contains_complex_structure(self, content_html):
        """
        Check if HTML content contains tables, lists, or other complex structures.
        
        Args:
            content_html (str): HTML content to analyze.
            
        Returns:
            bool: True if content contains tables/lists, False otherwise.
        """
        content_soup = BeautifulSoup(content_html, 'html.parser')
        
        # Check for tables
        if content_soup.find('table'):
            return True
        
        # Check for lists (ul, ol)
        if content_soup.find('ul') or content_soup.find('ol'):
            return True
        
        return False

    def __escape_hash_symbols(self, html_content):
        """
        Escape hash (#) symbols at the beginning of lines in HTML to prevent
        them from being treated as markdown headers.
        
        Args:
            html_content (str): HTML content.
            
        Returns:
            str: HTML with escaped hash symbols.
        """
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Process all text-containing elements
        for element in soup.find_all(['p', 'li', 'td', 'th', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            if element.string:
                text = element.string
                # If line starts with #, escape it
                if text.strip().startswith('#'):
                    element.string = text.replace('#', '\\#', 1)
        
        return str(soup)

    def __wrap_marked_sections_in_code_blocks(self, html, start_marker, end_marker):
        """
        Find content between markers and wrap appropriately:
        - Simple text/code → <pre><code> block
        - Tables/lists → Custom wrapper with preserved structure
        
        Args:
            html (str): The HTML content from Mammoth.
            start_marker (str): The start marker text.
            end_marker (str): The end marker text.
            
        Returns:
            str: HTML with marked sections wrapped appropriately.
        """
        import html as html_module
        
        # Mammoth escapes < and > to &lt; and &gt;, so we need to escape our markers too
        escaped_start = html_module.escape(start_marker)
        escaped_end = html_module.escape(end_marker)
        
        # Pattern to find content between HTML-escaped markers (including HTML tags)
        # The markers will be in separate <p> tags, and content in between
        pattern = re.compile(
            f'<p>{re.escape(escaped_start)}</p>(.*?)<p>{re.escape(escaped_end)}</p>',
            re.DOTALL
        )
        
        def replace_with_appropriate_wrapper(match):
            content = match.group(1)
            
            # Detect if content has complex structure (tables, lists)
            has_complex_structure = self.__contains_complex_structure(content)
            
            if has_complex_structure:
                # Preserve structure: keep HTML as-is, escape # symbols
                escaped_content = self.__escape_hash_symbols(content)
                # Wrap in a div with special class for potential custom handling
                return f'<div class="elitea-bordered-content">{escaped_content}</div>'
            else:
                # Simple text/code: extract as plain text and wrap in code block
                content_soup = BeautifulSoup(content, 'html.parser')
                
                # Extract text from each paragraph separately to preserve line breaks
                lines = []
                for element in content_soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    # Replace <br /> within paragraphs with newlines
                    for br in element.find_all('br'):
                        br.replace_with('\n')
                    text = element.get_text()
                    # Preserve leading whitespace (indentation), only strip trailing
                    lines.append(text.rstrip())
                
                # If no paragraphs found, just get all text
                if not lines:
                    content = content.replace('<br />', '\n').replace('<br/>', '\n').replace('<br>', '\n')
                    content_text = content_soup.get_text()
                    lines = [line.rstrip() for line in content_text.split('\n')]
                
                # Join lines, strip only leading/trailing empty lines
                content_text = '\n'.join(lines).strip()
                # Return as code block (need to HTML-escape the content)
                content_escaped = html_module.escape(content_text)
                return f'<pre><code>{content_escaped}</code></pre>'
        
        # Replace all marked sections with appropriate wrappers
        result_html = pattern.sub(replace_with_appropriate_wrapper, html)

        return result_html

    def __hoist_images_from_headings(self, html):
        """Relocate <img> tags that Mammoth placed inside heading elements.

        When an image is inserted into a heading paragraph in the source DOCX,
        Mammoth emits it as an inline <img> inside an <h1>..<h6>. markdownify
        renders headings as text-only and silently drops inline <img> tags.
        Move each such <img> into its own <p> immediately
        after the heading so it survives markdownify.
        The heading keeps its text.

        Args:
            html (str): The HTML content from Mammoth.

        Returns:
            str: serialized HTML with any heading-embedded images relocated.
                Every document is routed through this single serialization
                (see the return note below).
        """
        soup = BeautifulSoup(html, 'html.parser')
        heading_tags = ('h1', 'h2', 'h3', 'h4', 'h5', 'h6')
        for heading in soup.find_all(heading_tags):
            images = heading.find_all('img')
            if not images:
                continue
            # Whether the heading carries its own text besides the image(s).
            had_text = bool(heading.get_text(strip=True))
            anchor = heading
            for img in images:
                img.extract()
                paragraph = soup.new_tag('p')
                paragraph.append(img)
                anchor.insert_after(paragraph)
                anchor = paragraph
            if not had_text:
                # Image-only heading (e.g. a logo/banner used as a section
                # divider). Keep it as a header boundary — header-based chunking
                # relies on it — but give it a label rather than leaving an empty
                # '#'. Prefer the image's alt/filename, else a generic label.
                # Setting .string also clears any leftover empty inline wrappers.
                label = (images[0].get('alt') or '').strip() or 'Image'
                heading.string = label
        # Always return the serialized soup, even when nothing moved, so every
        # document follows one canonical HTML path. We already parsed `html`
        # above, and markdownify itself parses with BeautifulSoup, so this
        # round-trip is idempotent for mammoth's HTML and costs only a serialize
        # — but it removes the moved/not-moved divergence that would otherwise
        # be an input-specific, hard-to-reproduce difference.
        return str(soup)

    def load(self):
        """
        Loads and converts the Docx file to markdown format.

        Returns:
            List[Document]: A list containing a Documents with the markdown content
                          and metadata including the source file path.
        """
        result_content = self.get_content()
        return list(markdown_by_headers_chunker(iter([Document(page_content=result_content, metadata={'source': str(self.path)})]), config={'max_tokens':self.max_tokens}))

    def get_content(self):
        """Extracts and converts the content of the Docx file.

        When ``read_images_only`` is True and ``extracted_images_names`` is
        provided, bypasses mammoth entirely and returns a JSON dict of
        ``{filename: transcript}`` for the requested images only.

        When ``start_paragraph`` or ``end_paragraph`` is set, returns a fast
        structural slice via ``_read_paragraph_slice`` (no mammoth).
        """
        if self.read_images_only and self.extracted_images_names:
            return self._read_images_only()

        if self.path:
            with open(self.path, 'rb') as docx_file:
                return self._convert_docx_to_markdown(docx_file)
        elif self.file_content and self.file_name:
            docx_file = BytesIO(self.file_content)
            return self._convert_docx_to_markdown(docx_file)
        else:
            raise ValueError("Either 'path' or 'file_content' and 'file_name' must be provided.")

    def _read_images_only(self):
        """Extract and transcribe only the requested images, bypassing mammoth.

        Uses python-docx relationships to find images by filename and
        ``rel.target_part.blob`` for raw bytes.  Returns a JSON string of
        ``{filename: transcript_or_error}``.
        """
        from elitea_sdk.tools.utils.content_parser import image_processing_prompt

        if self.path:
            doc = DocxDocument(self.path)
        elif self.file_content:
            doc = DocxDocument(BytesIO(self.file_content) if isinstance(
                self.file_content, (bytes, bytearray)) else self.file_content)
        else:
            raise ValueError("Either 'path' or 'file_content' must be provided.")

        # Build filename → relationship mapping
        file_rels = {}
        for rel in doc.part.rels.values():
            try:
                tp = rel.target_part
            except Exception:  # pylint: disable=broad-except
                continue  # skip External relationships (hyperlinks etc.)
            if hasattr(tp, 'content_type') and tp.content_type.startswith('image/'):
                file_rels[tp.partname.filename] = rel

        results = {}
        for name in self.extracted_images_names:
            if name not in file_rels:
                results[name] = "Error: image not found in document"
                continue
            try:
                image_bytes = file_rels[name].target_part.blob
                if self.llm:
                    source = self.path or self.file_name or "docx"
                    transcript = perform_llm_prediction_for_image_bytes(
                        image_bytes, self.llm,
                        self.prompt if self.prompt else image_processing_prompt,
                        cache=self.image_cache,
                        image_name=source,
                    )
                else:
                    transcript = "Transcript is not available"
                results[name] = transcript
            except Exception as exc:
                results[name] = f"Error: {exc}"

        return json.dumps(results)

    def _convert_docx_to_markdown(self, docx_file):
        """Converts DOCX content to markdown with image handling.

        Detects bordered content and treats it as code blocks.  Pre-scans
        image references so that callbacks can include the original filename.
        """
        # Reset per-conversion image state so a reused loader instance does not
        # carry stale tokens or dedup hits across multiple conversions.
        self._reset_image_state()

        if hasattr(docx_file, 'seek'):
            docx_file.seek(0)

        # Step 0: Pre-scan image references for filename resolution
        pre_doc = DocxDocument(docx_file)
        self._scan_image_references(pre_doc)
        if hasattr(docx_file, 'seek'):
            docx_file.seek(0)

        # Step 1: Detect and mark bordered content
        marked_docx, start_marker, end_marker = self.__detect_and_mark_bordered_content(docx_file)
        
        # Step 2: Convert marked DOCX to HTML using Mammoth
        # Reset callback counter — bordered-content detection re-saves the
        # document which may change internal order, but the blip sequence is
        # preserved so the counter just needs to restart from 0.
        self._image_ref_index = 0
        if self.extract_images:
            result = convert_to_html(marked_docx, convert_image=mammoth.images.img_element(self.__handle_image))
        else:
            result = convert_to_html(marked_docx, convert_image=mammoth.images.img_element(self.__placeholder_image_handler))
        
        # Step 3: Wrap marked sections in <pre><code> tags
        html_with_code_blocks = self.__wrap_marked_sections_in_code_blocks(
            result.value, start_marker, end_marker
        )

        # Step 3.5: Hoist images out of headings so markdownify keeps them
        # (markdownify renders headings as text-only and drops inline <img>).
        html_with_code_blocks = self.__hoist_images_from_headings(html_with_code_blocks)

        # Step 4: Convert HTML to markdown
        content = markdownify(html_with_code_blocks, heading_style="ATX")
        
        # Step 5: Post-process markdown (for image transcripts, etc.)
        return self.__postprocess_original_md(content)
